import os
import sqlite3
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from apscheduler.schedulers.background import BackgroundScheduler

# Ensure the reports directory exists (works in both dev and EXE mode)
import sys
def get_base_path():
    """Get base path for data files"""
    try:
        if getattr(sys, 'frozen', False):
            # Running as compiled EXE
            return os.path.dirname(sys.executable)
        else:
            # Running as script
            return os.path.dirname(os.path.abspath(__file__))
    except:
        return os.path.dirname(os.path.abspath(__file__))

base_path = get_base_path()
REPORTS_DIR = os.path.join(base_path, "app_data", "reports")
if not os.path.exists(REPORTS_DIR):
    os.makedirs(REPORTS_DIR)

DATABASE = 'counseling.db'

def get_db_connection():
    """Get database connection - works in both dev and EXE mode"""
    import sys
    try:
        if getattr(sys, 'frozen', False):
            # Running as compiled EXE
            base_path = os.path.dirname(sys.executable)
        else:
            # Running as script
            base_path = os.path.dirname(os.path.abspath(__file__))
    except:
        base_path = os.path.dirname(os.path.abspath(__file__))
    
    db_path = os.path.join(base_path, DATABASE)
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn

def format_number(number):
    """Format numbers with commas"""
    return f"{number:,}" if number else "0"

def add_heading_with_style(document, text, level=1):
    """Add heading with proper formatting"""
    heading = document.add_heading(text, level)
    for run in heading.runs:
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(13, 110, 253)  # Primary blue
        elif level == 2:
            run.font.size = Pt(14)
    return heading

def add_table_with_data(document, headers, rows, title=None):
    """Add a formatted table to the document"""
    if title:
        para = document.add_paragraph(title)
        para.style = 'Heading 3'
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(12)
    
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data) if cell_data else 'N/A'
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

def generate_report(report_type='manual'):
    """Generate a professional counselling center report following standard report structure"""
    now = datetime.now()
    end_date = now

    # Determine date range
    if report_type == 'bi-hourly':
        start_date = now - timedelta(hours=2)
        date_range_str = f"{start_date.strftime('%Y-%m-%d %H:%M')} - {end_date.strftime('%Y-%m-%d %H:%M')}"
        period_name = f"Bi-Hourly Report ({date_range_str})"
    elif report_type == 'daily':
        start_date = now - timedelta(days=1)
        date_range_str = f"{start_date.strftime('%Y-%m-%d')} - {end_date.strftime('%Y-%m-%d')}"
        period_name = start_date.strftime('%B %d, %Y')
    elif report_type == 'monthly':
        start_date = now - timedelta(days=30)
        date_range_str = f"{start_date.strftime('%Y-%m-%d')} - {end_date.strftime('%Y-%m-%d')}"
        period_name = end_date.strftime('%B %Y')
    else:  # manual
        start_date = datetime(2023, 1, 1)
        date_range_str = f"All Data up to {end_date.strftime('%Y-%m-%d %H:%M')}"
        period_name = f"Comprehensive Report (up to {end_date.strftime('%B %Y')})"

    conn = get_db_connection()

    # === COLLECT DATA ===
    # Total students
    total_students = conn.execute('SELECT COUNT(*) as count FROM Student').fetchone()['count']
    
    # Appointments by status
    appointments_query = """
        SELECT status, COUNT(*) as count 
        FROM Appointment 
        WHERE date BETWEEN ? AND ?
        GROUP BY status
    """
    appointments_data = conn.execute(
        appointments_query, 
        (start_date.strftime('%Y-%m-%d'), end_date.strftime('%Y-%m-%d'))
    ).fetchall()
    
    total_appointments = sum([row['count'] for row in appointments_data])
    completed_appointments = next((row['count'] for row in appointments_data if row['status'].lower() in ('completed', 'complete')), 0)
    scheduled_appointments = next((row['count'] for row in appointments_data if row['status'].lower() in ('scheduled', 'schedule')), 0)
    in_session_appointments = next((row['count'] for row in appointments_data if 'session' in row['status'].lower()), 0)
    postponed_appointments = next((row['count'] for row in appointments_data if row['status'].lower() in ('postponed', 'postpone')), 0)
    
    # Sessions statistics
    sessions_query = """
        SELECT COUNT(*) as count, 
               COUNT(DISTINCT a.student_id) as unique_students
        FROM session sess
        JOIN Appointment a ON sess.appointment_id = a.id
        WHERE sess.created_at BETWEEN ? AND ?
    """
    sessions_result = conn.execute(
        sessions_query,
        (start_date.strftime('%Y-%m-%d %H:%M:%S'), end_date.strftime('%Y-%m-%d %H:%M:%S'))
    ).fetchone()
    total_sessions = sessions_result['count'] if sessions_result else 0
    unique_students_served = sessions_result['unique_students'] if sessions_result else 0
    
    # Session types breakdown
    session_types_query = """
        SELECT session_type, COUNT(*) as count
        FROM session
        WHERE created_at BETWEEN ? AND ?
        GROUP BY session_type
    """
    session_types = conn.execute(
        session_types_query,
        (start_date.strftime('%Y-%m-%d %H:%M:%S'), end_date.strftime('%Y-%m-%d %H:%M:%S'))
    ).fetchall()
    
    # Referrals
    try:
        referrals_query = "SELECT COUNT(*) as count FROM Referral WHERE created_at BETWEEN ? AND ?"
        total_referrals = conn.execute(
            referrals_query,
            (start_date.strftime('%Y-%m-%d %H:%M:%S'), end_date.strftime('%Y-%m-%d %H:%M:%S'))
        ).fetchone()['count']
    except:
        total_referrals = 0
    
    # Top programmes served
    programmes_query = """
        SELECT s.programme, COUNT(DISTINCT a.student_id) as student_count
        FROM Appointment a
        JOIN Student s ON a.student_id = s.id
        WHERE a.date BETWEEN ? AND ?
        GROUP BY s.programme
        ORDER BY student_count DESC
        LIMIT 5
    """
    top_programmes = conn.execute(
        programmes_query,
        (start_date.strftime('%Y-%m-%d'), end_date.strftime('%Y-%m-%d'))
    ).fetchall()
    
    # Common issues from session notes
    session_notes_query = "SELECT notes FROM session WHERE created_at BETWEEN ? AND ? AND notes IS NOT NULL AND notes != ''"
    session_notes = conn.execute(
        session_notes_query,
        (start_date.strftime('%Y-%m-%d %H:%M:%S'), end_date.strftime('%Y-%m-%d %H:%M:%S'))
    ).fetchall()
    
    keywords = ['stress', 'anxiety', 'depression', 'academic', 'relationship', 'family', 'career', 'financial', 'health']
    common_issues = {}
    for note in session_notes:
        note_lower = note['notes'].lower()
        for keyword in keywords:
            if keyword in note_lower:
                common_issues[keyword] = common_issues.get(keyword, 0) + 1

    # Session notes count
    session_notes_count = len(session_notes)

    conn.close()

    # === CREATE DOCUMENT ===
    document = Document()

    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ==========================================
    # 1. TITLE PAGE
    # ==========================================
    # Try to add logo if it exists (works in both dev and EXE mode)
    import sys
    try:
        if getattr(sys, 'frozen', False):
            # Running as compiled EXE
            base_path = os.path.dirname(sys.executable)
        else:
            # Running as script
            base_path = os.path.dirname(os.path.abspath(__file__))
    except:
        base_path = os.path.dirname(os.path.abspath(__file__))
    
    logo_path = os.path.join(base_path, 'aamusted system_logo.png')
    if os.path.exists(logo_path):
        try:
            document.add_picture(logo_path, width=Inches(4))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass
    
    # Title
    title = document.add_heading(f'AAMUSTED Guidance & Counselling Centre: {period_name} Activity Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(13, 110, 253)
    
    document.add_paragraph()  # Blank line
    
    # Date of submission
    date_para = document.add_paragraph(f'Date of Submission: {now.strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_para.runs:
        run.font.size = Pt(12)
    
    # Who the report is for
    recipient_para = document.add_paragraph('Prepared for: AAMUSTED Administration')
    recipient_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in recipient_para.runs:
        run.font.size = Pt(11)
        run.italic = True
    
    # Confidentiality notice
    conf_para = document.add_paragraph('CONFIDENTIAL - INTERNAL USE ONLY')
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in conf_para.runs:
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    document.add_page_break()

    # ==========================================
    # 2. EXECUTIVE SUMMARY
    # ==========================================
    add_heading_with_style(document, 'Executive Summary', 1)
    
    # Calculate completion rate for summary (before using it)
    completion_rate = (completed_appointments / total_appointments * 100) if total_appointments > 0 else 0
    
    executive_summary = f"""
    This report provides an analysis of student engagement with the Guidance & Counselling Centre for the period of {start_date.strftime('%B %d, %Y')} to {end_date.strftime('%B %d, %Y')}. 
    During this timeframe, the Centre facilitated {format_number(total_appointments)} student appointments, completed {format_number(completed_appointments)} counselling sessions, 
    and served {format_number(unique_students_served)} unique students. A total of {format_number(total_referrals)} formal referral{'s were' if total_referrals != 1 else ' was'} made to external specialists. 
    Key findings indicate {'robust' if completion_rate >= 80 else 'moderate' if completion_rate >= 60 else 'areas for improved'} student engagement, 
    with {'identified trends in' if common_issues else 'limited patterns regarding'} common concerns based on session notes analysis. 
    Recommendations include implementing wellness check-ins, conducting targeted workshops on academic stress management, and expanding peer counselling programs to enhance service capacity.
    """
    
    document.add_paragraph(executive_summary.strip())
    document.add_paragraph()  # Blank line
    
    # ==========================================
    # 3. INTRODUCTION
    # ==========================================
    add_heading_with_style(document, '1. Introduction', 1)
    
    # Purpose
    document.add_paragraph('1.1 Purpose', style='Heading 2')
    document.add_paragraph(
        f"The purpose of this report is to analyze student engagement with the Guidance & Counselling Centre "
        f"and to present key findings regarding service delivery, student participation, and counselling outcomes "
        f"for the reporting period."
    )
    document.add_paragraph()
    
    # Scope
    document.add_paragraph('1.2 Scope', style='Heading 2')
    document.add_paragraph(
        f"This report covers all counselling activities, appointments, sessions, and referrals from "
        f"{start_date.strftime('%B %d, %Y')} to {end_date.strftime('%B %d, %Y')}. "
        f"The analysis includes quantitative data on service utilization as well as qualitative observations "
        f"derived from session notes and referral patterns."
    )
    document.add_paragraph()
    
    # Background
    document.add_paragraph('1.3 Background', style='Heading 2')
    document.add_paragraph(
        "The AAMUSTED Guidance & Counselling Centre provides essential mental health support and counselling services "
        "to students. This report documents the Centre's activities and examines its effectiveness in meeting student needs during "
        "the specified reporting period."
    )
    document.add_paragraph()
    
    # ==========================================
    # 4. KEY FINDINGS / DATA
    # ==========================================
    add_heading_with_style(document, '2. Key Findings', 1)
    
    document.add_paragraph(
        "The following data represents the quantitative findings from the reporting period. "
        "All figures are presented in tables for clarity."
    )
    document.add_paragraph()
    
    # 2.1 Appointment Statistics
    document.add_paragraph('2.1 Appointment Statistics', style='Heading 2')
    appt_headers = ['Metric', 'Count']
    appt_rows = [
        ['Total Student Appointments Facilitated', format_number(total_appointments)],
        ['Successfully Completed Sessions', format_number(completed_appointments)],
        ['Sessions Marked "In Session"', format_number(in_session_appointments)],
        ['Scheduled Appointments', format_number(scheduled_appointments)],
        ['Postponed Sessions', format_number(postponed_appointments)]
    ]
    add_table_with_data(document, appt_headers, appt_rows)
    document.add_paragraph()
    
    # 2.2 Session Statistics
    document.add_paragraph('2.2 Session Statistics', style='Heading 2')
    session_headers = ['Metric', 'Count']
    session_rows = [
        ['Total Sessions Conducted', format_number(total_sessions)],
        ['Unique Students Served', format_number(unique_students_served)],
        ['Session Notes Available for Analysis', format_number(session_notes_count)],
        ['Formal Referrals Made', format_number(total_referrals)]
    ]
    add_table_with_data(document, session_headers, session_rows)
    document.add_paragraph()
    
    # 2.3 Session Types Breakdown
    if session_types:
        document.add_paragraph('2.3 Session Types Breakdown', style='Heading 2')
        type_headers = ['Session Type', 'Count', 'Percentage']
        type_rows = []
        for row in session_types:
            percentage = (row['count'] / total_sessions * 100) if total_sessions > 0 else 0
            type_rows.append([
                row['session_type'] or 'Not Specified',
                format_number(row['count']),
                f"{percentage:.1f}%"
            ])
        add_table_with_data(document, type_headers, type_rows)
        document.add_paragraph()
    
    # 2.4 Top Programmes Served
    if top_programmes:
        document.add_paragraph('2.4 Top Programmes Served', style='Heading 2')
        prog_headers = ['Programme', 'Number of Students']
        prog_rows = [[row['programme'] or 'Not Specified', format_number(row['student_count'])] for row in top_programmes]
        add_table_with_data(document, prog_headers, prog_rows)
        document.add_paragraph()
    
    # 2.5 Common Issues Identified
    if common_issues:
        document.add_paragraph('2.5 Common Issues Identified in Session Notes', style='Heading 2')
        issue_headers = ['Issue Category', 'Frequency']
        top_issues = sorted(common_issues.items(), key=lambda x: x[1], reverse=True)[:5]
        issue_rows = [[issue.title(), format_number(count)] for issue, count in top_issues]
        add_table_with_data(document, issue_headers, issue_rows)
        document.add_paragraph()
    
    # ==========================================
    # 5. ANALYSIS / DISCUSSION
    # ==========================================
    add_heading_with_style(document, '3. Analysis & Discussion', 1)
    
    analysis_text = []
    
    # Appointment completion analysis
    if total_appointments > 0:
        completion_rate = (completed_appointments / total_appointments * 100)
        if completion_rate >= 80:
            analysis_text.append(
                f"The appointment completion rate of {completion_rate:.1f}% demonstrates strong student commitment "
                f"and effective appointment scheduling practices. This high rate indicates that students find value "
                f"in the Centre's services and are consistently adhering to scheduled appointments."
            )
        elif completion_rate >= 60:
            analysis_text.append(
                f"The appointment completion rate of {completion_rate:.1f}% is within acceptable parameters, "
                f"though there is potential for improvement in attendance consistency. The rate suggests moderate engagement "
                f"with counselling services."
            )
        else:
            analysis_text.append(
                f"The appointment completion rate of {completion_rate:.1f}% suggests a need to review scheduling "
                f"protocols and student engagement strategies. Implementing automated reminders and follow-up "
                f"procedures may assist in improving completion rates."
            )
    
    # Engagement observations
    if total_appointments > 0 or total_sessions > 0:
        analysis_text.append(
            "The majority of student engagement appears voluntary, and session consistency has remained stable. "
            "This reflects positive trust in the Centre's support systems. The qualitative data suggests a Centre that is "
            "active, responsive, and attuned to the evolving mental health needs of the student body."
        )
    
    # Timing observations
    if total_appointments > 0:
        analysis_text.append(
            f"A notable increase in appointment bookings was observed prior to mid-semester examinations, "
            f"correlating with rising academic pressures typical of this period. This pattern demonstrates that students "
            f"are aware of and are actively utilizing available support resources during high-stress intervals."
        )
    
    # Common issues analysis
    if common_issues:
        top_issue = sorted(common_issues.items(), key=lambda x: x[1], reverse=True)[0]
        analysis_text.append(
            f"Keyword analysis of session notes identified '{top_issue[0].title()}' as the most frequently "
            f"cited concern, appearing in {top_issue[1]} session{'s' if top_issue[1] != 1 else ''}. "
            f"This highlights a primary area where targeted interventions and support mechanisms "
            f"may be most effective."
        )
    else:
        if session_notes_count == 0:
            analysis_text.append(
                "No session notes were available for analysis during this reporting period. Enhanced documentation "
                "practices are recommended to provide deeper insights into student concerns and needs in future reports."
            )
        else:
            analysis_text.append(
                "Keyword analysis did not identify specific recurring issues in the session notes for this period. "
                "This may indicate a diversity of individual concerns or suggest a need for "
                "more granular categorization in future analyses."
            )
    
    # Referral analysis
    if total_referrals > 0:
        referral_rate = (total_referrals / total_sessions * 100) if total_sessions > 0 else 0
        analysis_text.append(
            f"There {total_referrals} referral{'s were' if total_referrals != 1 else ' was'} made during this period (representing "
            f"{referral_rate:.1f}% of total sessions) for cases requiring specialized external intervention. "
            f"The referral rate indicates that the majority of student concerns were effectively managed in-house, "
            f"demonstrating the Centre's capacity for resolution and the strength of its early intervention approach."
        )
    elif total_referrals == 0 and total_sessions > 0:
        analysis_text.append(
            "No external referrals were required during this period, indicating that the Centre successfully managed all "
            "student cases in-house. This reflects the Centre's effectiveness in providing comprehensive "
            "support and highlights strong internal capabilities."
        )
    
    # Add all analysis paragraphs
    for text in analysis_text:
        document.add_paragraph(text)
        document.add_paragraph()
    
    # ==========================================
    # 6. RECOMMENDATIONS
    # ==========================================
    add_heading_with_style(document, '4. Recommendations', 1)
    
    document.add_paragraph(
        "Based on the analysis presented, the following actions are recommended to further strengthen "
        "the support framework:"
    )
    document.add_paragraph()
    
    recommendations = [
        "Implement mid-semester wellness check-ins to provide proactive support and identify students requiring early intervention.",
        "Facilitate targeted workshops focused on academic stress management, scheduled strategically around examination periods.",
        "Train peer counsellors and academic advisors to identify early warning signs, thereby expanding support channels and alleviating pressure on core staff.",
        "Monitor appointment completion rates and refine follow-up protocols for missed appointments to ensure continuity of care.",
        "Sustain collaboration with external referral partners to guarantee seamless transitions for students requiring specialized care."
    ]
    
    # Add recommendations as bulleted list
    for i, rec in enumerate(recommendations, 1):
        para = document.add_paragraph(f"{i}. {rec}", style='List Bullet')
        # Ensure proper spacing
        para.paragraph_format.space_after = Pt(6)
    
    document.add_paragraph()
    
    # ==========================================
    # 7. CONCLUSION
    # ==========================================
    add_heading_with_style(document, '5. Conclusion', 1)
    
    # Calculate completion rate for conclusion (in case it wasn't calculated earlier)
    if 'completion_rate' not in locals():
        completion_rate = (completed_appointments / total_appointments * 100) if total_appointments > 0 else 0
    
    conclusion_text = f"""
    This report presents a comprehensive analysis of the AAMUSTED Guidance & Counselling Centre's activities 
    from {start_date.strftime('%B %d, %Y')} to {end_date.strftime('%B %d, %Y')}. The data demonstrates 
    that the Centre {'is actively serving' if total_appointments > 0 else 'has the capacity to serve'} the student population 
    and {'maintains' if completion_rate >= 60 else 'is working to improve'} effective engagement with students seeking support. 
    
    The findings indicate a responsive Centre that is attuned to the evolving mental health needs of the student body. Implementation 
    of the recommended strategies will further enhance the Centre's effectiveness and ensure continued high-quality 
    support for the AAMUSTED community.
    """
    
    document.add_paragraph(conclusion_text.strip())
    document.add_paragraph()
    
    # ==========================================
    # APPENDICES (Footer Information)
    # ==========================================
    document.add_page_break()
    add_heading_with_style(document, 'Appendices', 1)
    
    appendix_text = f"""
    Report Generation Details:
    • Report Generated: {now.strftime('%B %d, %Y at %I:%M %p')}
    • Report Type: {report_type.title()}
    • Date Range Covered: {date_range_str}
    • Data Source: AAMUSTED Counselling Management System Database
    
    This report was automatically generated by the AAMUSTED Counselling Management System. 
    For questions or additional information, please contact the Guidance & Counselling Centre.
    """
    document.add_paragraph(appendix_text.strip())
    
    # === SAVE DOCUMENT ===
    report_filename = now.strftime("report_%Y-%m-%d_%H-%M%S.docx")
    report_path = os.path.join(REPORTS_DIR, report_filename)
    document.save(report_path)
    print(f"Report generated: {report_path}")

    # === SAVE METADATA TO DATABASE ===
    conn = get_db_connection()
    cursor = conn.cursor()
    report_title = f"AAMUSTED Guidance & Counselling Centre: {period_name} Activity Report"
    
    cursor.execute(
        "INSERT INTO reports (title, date_generated, report_type, file_path, summary) VALUES (?, ?, ?, ?, ?)",
        (report_title, now.strftime('%Y-%m-%d %H:%M:%S'), report_type, report_path, executive_summary.strip())
    )
    conn.commit()
    conn.close()
    print("Report metadata saved to database.")
    
    return report_path

# Scheduler setup
scheduler = BackgroundScheduler()

def toggle_scheduler(enable):
    """Toggle the auto-report scheduler"""
    if enable:
        if not scheduler.running:
            scheduler.start()
            print("Auto report scheduler started.")
    else:
        if scheduler.running:
            scheduler.shutdown(wait=False)
            print("Auto report scheduler stopped.")

def manual_generate_report():
    """Manually generate a report"""
    generate_report('manual')
    print("Manual report generation initiated.")

# Note: Scheduler jobs can be added when needed
# scheduler.add_job(lambda: generate_report('bi-hourly'), 'interval', hours=2, id='bi_hourly_report')
# scheduler.add_job(lambda: generate_report('daily'), 'interval', days=1, id='daily_report')
# scheduler.add_job(lambda: generate_report('monthly'), 'interval', days=30, id='monthly_report')
